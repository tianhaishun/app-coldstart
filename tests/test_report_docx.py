"""Word 性能报告导出测试。

覆盖两类回归：
1. generate_report 冒烟：返回字节可被 python-docx 打开、标题/设备内容正确、
   iOS 首次调整生效。
2. 导出端点头部编码：中文标题曾把 Content-Disposition 拼进 latin-1 编码的
   响应头，starlette 构造 Response 时抛 UnicodeEncodeError → 100% 500。
"""
import io

import pytest
from docx import Document

from server import ReportExportReq, export_report_docx


def _device(serial, firsts_ms, seconds_ms):
    """构造一台设备的报告数据（time 单位毫秒）。"""
    return {
        "serial": serial,
        "model": f"Model-{serial}",
        "label": f"设备 {serial}",
        "brand": "TestBrand",
        "osVersion": "Android 15",
        "firsts": [{"time": t, "date": "2026-08-13 10:00:00", "abnormal": False} for t in firsts_ms],
        "seconds": [{"time": t, "date": "2026-08-13 10:01:00", "abnormal": False} for t in seconds_ms],
    }


def _base_data(**over):
    data = {
        "title": "20260813-测试应用 冷启动测速报告",
        "testDate": "2026-08-13 10:00",
        "appName": "com.example.app",
        "platform": "gp",
        "platformLabel": "GP（安卓）",
        "launchMode": "模拟点击图标",
        "launchModeDesc": "模拟点击图标",
        "plannedRounds": 5,
        "totalSec": 120.5,
        "success": True,
        "devices": [_device("DEV_A", [2000, 2400, 2600, 3000, 2800], [1500, 1600, 1700])],
        "auditLog": ["[10:00:00] 开始测速", "[10:02:00] 完成"],
    }
    data.update(over)
    return data


def _doc_text(doc):
    """提取 docx 全部文本（段落 + 表格单元格）。"""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def test_generate_report_smoke():
    """generate_report 返回可用 docx，标题与关键内容齐全。"""
    from report_docx import generate_report

    data = _base_data()
    blob = generate_report(data)
    assert isinstance(blob, bytes) and blob[:2] == b"PK"  # docx 是 zip 容器
    doc = Document(io.BytesIO(blob))
    text = _doc_text(doc)
    assert data["title"] in text
    assert "com.example.app" in text
    assert "设备 DEV_A" in text
    # 明细：异常标记/审计日志不应丢
    assert "开始测速" in text


def test_generate_report_ios_first_adjust():
    """platform=ios 时首次均值减 iosFirstAdjustSec（默认 1.0s）。

    firsts=[5230,6410,7120,8340,9990]ms → 秒 [5.23,6.41,7.12,8.34,9.99]
    剔 5.23/9.99 → 均值 (6.41+7.12+8.34)/3=7.29 → 调整后 6.29。
    gp 不调整，应出现 7.29 而非 6.29。
    """
    from report_docx import generate_report

    firsts_ms = [5230, 6410, 7120, 8340, 9990]
    seconds_ms = [1234, 2567, 3987]  # 均值 2.60（剔除后 2.567）
    ios_text = _doc_text(Document(io.BytesIO(generate_report(
        _base_data(platform="ios", platformLabel="iOS",
                   devices=[_device("DEV_I", firsts_ms, seconds_ms)])))))
    gp_text = _doc_text(Document(io.BytesIO(generate_report(
        _base_data(devices=[_device("DEV_G", firsts_ms, seconds_ms)])))))

    assert "6.29" in ios_text, "iOS 首次均值应显示 7.29-1.0=6.29"
    assert "7.29" not in ios_text
    assert "7.29" in gp_text, "gp 首次均值不应被调整"
    assert "6.29" not in gp_text


def test_export_report_docx_chinese_title_no_500():
    """中文标题曾 100% 触发 latin-1 UnicodeEncodeError → 500（实测回归）。"""
    req = ReportExportReq(**_base_data(title="20260813-某应用 冷启动测速报告"))
    resp = export_report_docx(req)  # 修复前此处抛 HTTPException(500)
    assert resp.status_code == 200
    cd = resp.headers["content-disposition"]
    # filename= 必须是纯 ASCII（latin-1 可编码）；中文名走 filename*=UTF-8''
    assert 'filename="' in cd
    filename_part = cd.split("filename=", 1)[1].split(";", 1)[0].strip('"')
    assert filename_part.encode("latin-1").decode("latin-1") == filename_part  # 纯 ASCII
    assert "filename*=UTF-8''" in cd
    assert "%E5%86%B7" in cd  # 「冷」的 UTF-8 percent 编码
    # 返回体仍是合法 docx
    doc = Document(io.BytesIO(resp.body))
    assert "某应用" in _doc_text(doc)


def test_export_report_docx_ascii_title():
    """纯 ASCII 标题：filename= 即标题本身，filename* 仍附带。"""
    req = ReportExportReq(**_base_data(title="MyReport-2026"))
    resp = export_report_docx(req)
    assert resp.status_code == 200
    cd = resp.headers["content-disposition"]
    assert 'filename="MyReport-2026.docx"' in cd


def test_export_report_docx_special_chars_title():
    """标题含引号/斜杠等特殊字符时只影响文件名，不影响生成。"""
    req = ReportExportReq(**_base_data(title='20260813-应用"测试"/报告'))
    resp = export_report_docx(req)
    assert resp.status_code == 200
    filename_part = resp.headers["content-disposition"].split("filename=", 1)[1].split(";", 1)[0].strip('"')
    assert '"' not in filename_part and "/" not in filename_part
