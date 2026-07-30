#!/usr/bin/env python3
# ──────────────────────────────────────────────────────────────────────────────
# 生成 Word 发布说明文档（.docx），可一键导入飞书 / 钉钉 / 语雀等在线文档。
#
# 数据来源：
#   - package.json          → 版本号、作者、描述
#   - CHANGELOG.md          → 更新内容（按版本段解析）
#   - release/*.exe         → 安装包文件名 + 体积
#
# 用法：
#   python scripts/gen-release-doc.py                # 默认输出到 publish/
#   python scripts/gen-release-doc.py -o /path/to/   # 指定输出目录
#
# 依赖：python-docx（pip install python-docx）
# ──────────────────────────────────────────────────────────────────────────────

import json
import re
import sys
import os
import argparse
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("[FAIL] 缺少 python-docx，请先安装：pip install python-docx")
    sys.exit(1)

# ── 路径 ──────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent
PKG_PATH = ROOT / "package.json"
CHANGELOG_PATH = ROOT / "CHANGELOG.md"
RELEASE_DIR = ROOT / "release"


# ── CHANGELOG 解析 ────────────────────────────────────────────────────────────

def parse_changelog(md_text):
    """
    解析 CHANGELOG.md，返回 [(version_label, sections_dict), ...]
    sections_dict = {"新增": [...lines], "修复": [...lines], "优化": [...lines], ...}

    支持两种标题格式：
      ## v2.0.0 — 2026-07-30         （语义版本）
      ## 2026-07-29 · 描述            （旧日期格式）
    """
    versions = []
    current_version = None
    current_category = None  # "新增" / "修复" / ...
    current_lines = {}

    for line in md_text.splitlines():
        # 版本标题
        ver_match = re.match(r"^##\s+(.+)$", line)
        if ver_match:
            if current_version:
                versions.append((current_version, current_lines))
            current_version = ver_match.group(1).strip()
            current_category = None
            current_lines = {}
            continue

        # 跳过 "---" 分隔线和 "待办" 等
        if line.strip() in ("---",) or (current_version and "待办" in (current_version or "")):
            # 如果遇到待办段，不收集
            if current_version and "待办" in current_version:
                continue

        if not current_version:
            continue
        # 跳过待办段
        if "待办" in current_version or "未发布" in current_version.lower() and "unreleased" in current_version.lower():
            continue

        # 分类标题 ### 新增 / ### 修复 / ...
        cat_match = re.match(r"^###\s+(.+)$", line)
        if cat_match:
            current_category = cat_match.group(1).strip()
            if current_category not in current_lines:
                current_lines[current_category] = []
            continue

        # 内容行（列表项或普通段落）
        stripped = line.strip()
        if stripped and not stripped.startswith("#"):
            if current_category:
                current_lines[current_category].append(stripped)
            else:
                # 无分类标题的普通内容，归入 "概要"
                if "概要" not in current_lines:
                    current_lines["概要"] = []
                current_lines["概要"].append(stripped)

    if current_version:
        versions.append((current_version, current_lines))

    return versions


def format_version_label(label):
    """清理版本标签，去掉多余的标点"""
    return label.rstrip("—-·").strip()


# ── 文档样式辅助 ──────────────────────────────────────────────────────────────

BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)      # 深色标题
BRAND_BLUE = RGBColor(0x0D, 0x6E, 0xFD)       # 品牌蓝
BRAND_ACCENT = RGBColor(0xFA, 0xB2, 0x83)      # OC-2 暖桃
TEXT_GRAY = RGBColor(0x44, 0x44, 0x44)
TEXT_LIGHT = RGBColor(0x88, 0x88, 0x88)

# 分类标题的中文映射 + 图标
CATEGORY_ICONS = {
    "新增": "✨",
    "修复": "🐛",
    "优化": "⚡",
    "维护": "🔧",
    "接口": "🔌",
    "安全": "🔒",
    "概要": "📋",
}


def set_cell_shading(cell, color_hex):
    """设置单元格背景色"""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.find(qn("w:shd"))
    if shading_elm is None:
        from lxml import etree
        shading_elm = etree.SubElement(shading, qn("w:shd"))
    shading_elm.set(qn("w:fill"), color_hex)


def add_styled_heading(doc, text, level=1, color=BRAND_DARK):
    """添加带颜色的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = color
        run.font.name = "Microsoft YaHei"
        # 中文字体
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return heading


def add_bullet(doc, text, bold_prefix=None):
    """
    添加列表项。
    如果 bold_prefix 不为 None，则前半段加粗（粗体标题：描述）。
    """
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run_bold = para.add_run(bold_prefix)
        run_bold.bold = True
        run_bold.font.name = "Microsoft YaHei"
        run_bold._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run_bold.font.size = Pt(11)
        run_rest = para.add_run(text)
    else:
        run = para.add_run(text)
        run.font.size = Pt(11)

    run_rest = para.runs[-1]
    run_rest.font.name = "Microsoft YaHei"
    run_rest._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return para


def parse_bullet_line(line):
    """
    从 CHANGELOG 行中提取粗体前缀和正文。
    格式：- **标题**：描述内容
    返回 (prefix, rest)
    """
    # 去掉前导 - 或 *
    line = re.sub(r"^[-*]\s*", "", line)
    # 匹配 **粗体**：后续
    m = re.match(r"\*\*(.+?)\*\*[：:]\s*(.*)", line)
    if m:
        return (m.group(1), m.group(2))
    return (None, line)


# ── 文档生成 ──────────────────────────────────────────────────────────────────

def generate_doc(pkg, changelog_versions, installer_path, output_path):
    """
    生成完整的 .docx 发布说明。

    结构：
      1. 标题 + 版本信息表
      2. 应用简介
      3. 安装方式
      4. 更新内容（从 CHANGELOG 解析）
      5. 系统要求
      6. 已知限制
    """
    version = pkg.get("version", "0.0.0")
    author = pkg.get("author", "")
    product_name = pkg.get("build", {}).get("productName", "AppColdStart")
    today = date.today().strftime("%Y-%m-%d")

    doc = Document()

    # ── 全局默认字体 ──
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ── 页边距 ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # 1. 标题
    # ══════════════════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("App 冷启测速")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_DARK
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"v{version}  发布说明")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_BLUE
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph()  # 空行

    # ── 版本信息表 ──
    installer_name = Path(installer_path).name if installer_path else f"{product_name}-{version}-setup.exe"
    installer_size = ""
    if installer_path and Path(installer_path).exists():
        size_mb = Path(installer_path).stat().st_size / (1024 * 1024)
        installer_size = f"{size_mb:.0f} MB"

    info_rows = [
        ("版本号", f"v{version}"),
        ("发布日期", today),
        ("作者", author),
        ("安装包", installer_name),
        ("包大小", installer_size or "—"),
        ("适用平台", "Windows 10 / 11（64 位）"),
    ]

    table = doc.add_table(rows=len(info_rows), cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (key, val) in enumerate(info_rows):
        row = table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        # 左列加粗 + 背景
        for para in row.cells[0].paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(row.cells[0], "F0F0F5")
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                r.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. 简介
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "一、应用简介", level=1)

    intro = doc.add_paragraph()
    intro.add_run(
        "App 冷启测速（Cold Start Profiler）是一款 Android / iOS 双平台冷启动性能测试工具。"
        "支持全自动测速循环（卸装→测首次→杀进程→测二次），"
        "模板比对毫秒级自动停表，scrcpy 实时镜像与录屏，项目化持久化管理测试配置。"
    ).font.size = Pt(11)

    capabilities = [
        "全自动测速：卸装→首次冷启动→杀进程→二次冷启动，循环 N 轮，自动去极值统计",
        "模板比对停表：点选启动成功元素，运行时区域匹配，精度 ±50ms",
        "scrcpy 实时镜像：独立置顶窗口 30fps 投屏，支持后台录屏",
        "iOS 支持：USB 连接 iPhone，截图 + 模板比对 + 冷启动计时",
        "项目持久化：启动模板 / 跳过按钮 / 包名按项目分开存储",
        "内置工具链：ADB + scrcpy + libimobiledevice，不污染系统 PATH",
    ]
    for cap in capabilities:
        add_bullet(doc, cap)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. 安装方式
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "二、安装方式", level=1)

    add_styled_heading(doc, "桌面客户端（推荐）", level=2, color=BRAND_BLUE)
    steps = [
        "双击安装包 AppColdStart-2.0.0-setup.exe，按向导完成安装",
        "安装完成后，桌面 / 开始菜单出现 AppColdStart 图标",
        "双击图标启动，首次运行自动创建 Python 环境 + 安装依赖（约 1-3 分钟）",
        "将手机用 USB 连接电脑，开启 USB 调试（Android）或信任电脑（iOS）",
        "在应用中选中设备，开始测速",
    ]
    for i, step in enumerate(steps, 1):
        para = doc.add_paragraph(style="List Number")
        run = para.add_run(step)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_styled_heading(doc, "浏览器模式（备选）", level=2, color=BRAND_BLUE)
    add_bullet(doc, "双击 Start.bat，首次自动建虚拟环境 + 装依赖，启动后自动打开浏览器")
    add_bullet(doc, "浏览器模式无 scrcpy 镜像 / 录屏 / iOS 支持，仅 Android 截图轮询", None)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. 更新内容
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "三、更新内容", level=1)

    # 只展示最近 3 个版本段，避免文档过长
    max_versions = 3
    for i, (ver_label, sections) in enumerate(changelog_versions[:max_versions]):
        label = format_version_label(ver_label)
        # 跳过待办/未发布段
        if "待办" in label or "未发布" in label:
            continue

        add_styled_heading(doc, label, level=2, color=BRAND_DARK)

        for cat_name, lines in sections.items():
            if not lines:
                continue
            icon = CATEGORY_ICONS.get(cat_name, "•")
            # 分类小标题
            cat_para = doc.add_paragraph()
            run = cat_para.add_run(f"{icon} {cat_name}")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = BRAND_ACCENT
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

            for line in lines:
                prefix, rest = parse_bullet_line(line)
                add_bullet(doc, rest, bold_prefix=(f"{prefix}：" if prefix else None))

        if i < min(len(changelog_versions), max_versions) - 1:
            doc.add_paragraph()  # 版本间空行

    # 如果有更多历史版本，提示一下
    if len(changelog_versions) > max_versions:
        note = doc.add_paragraph()
        run = note.add_run(f"（更多历史版本详见 CHANGELOG.md，共 {len(changelog_versions)} 个版本段）")
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_LIGHT
        run.italic = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. 系统要求
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "四、系统要求", level=1)

    req_table = doc.add_table(rows=5, cols=2, style="Light List Accent 1")
    req_data = [
        ("操作系统", "Windows 10 / 11（64 位）"),
        ("Python", "3.10+（首次启动自动配置，需提前装好并加入 PATH）"),
        ("Android 设备", "USB 调试模式，数据线连接"),
        ("iOS 设备", "iTunes 或 AMDS 驱动，iPhone 解锁并信任电脑"),
        ("磁盘空间", "约 500 MB（含 Python 虚拟环境 + OCR 模型）"),
    ]
    for i, (key, val) in enumerate(req_data):
        row = req_table.rows[i]
        row.cells[0].text = key
        row.cells[1].text = val
        for para in row.cells[0].paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
        for para in row.cells[1].paragraphs:
            for r in para.runs:
                r.font.size = Pt(10)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. 已知限制
    # ══════════════════════════════════════════════════════════════════════════
    add_styled_heading(doc, "五、已知限制", level=1)

    limitations = [
        "计时起点漏掉 adb input tap 巧行链执行时间（约 150-250ms），横向对比不受影响",
        "iOS 非越狱设备不支持模拟点击，跳过弹窗需手动操作",
        "iOS 杀进程需手动在 App 切换器中上滑关闭",
        "scrcpy 镜像仅支持 Android，iOS 无投屏功能",
    ]
    for lim in limitations:
        add_bullet(doc, lim)

    doc.add_paragraph()
    doc.add_paragraph()

    # ── 页脚 ──
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"— App 冷启测速 v{version} · {today} · 作者：{author} —")
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT

    # ── 保存 ──
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] Word 文档已生成: {output_path}")
    return output_path


# ── 主入口 ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="生成 Word 发布说明文档")
    parser.add_argument("-o", "--output", default=None, help="输出目录（默认 publish/）")
    parser.add_argument("--version", default=None, help="指定版本号（默认读 package.json）")
    args = parser.parse_args()

    # 读取 package.json
    with open(PKG_PATH, "r", encoding="utf-8") as f:
        pkg = json.load(f)

    version = args.version or pkg.get("version", "0.0.0")

    # 查找安装包
    installer_path = None
    if RELEASE_DIR.exists():
        for f in RELEASE_DIR.glob("*-setup.exe"):
            installer_path = f
            break

    # 解析 CHANGELOG
    changelog_text = ""
    if CHANGELOG_PATH.exists():
        changelog_text = CHANGELOG_PATH.read_text(encoding="utf-8")
    versions = parse_changelog(changelog_text)

    if not versions:
        print("[WARN] CHANGELOG.md 未解析到版本段，文档将不含更新内容。")

    # 输出路径
    output_dir = Path(args.output) if args.output else (ROOT / "publish" / f"v{version}")
    installer_name = Path(installer_path).stem if installer_path else f"AppColdStart-{version}-setup"
    output_path = output_dir / f"发布说明-v{version}.docx"

    generate_doc(pkg, versions, installer_path, output_path)

    # 同时把安装包复制到 publish 目录
    if installer_path and installer_path.exists():
        import shutil
        dest = output_dir / installer_path.name
        if not dest.exists():
            shutil.copy2(str(installer_path), str(dest))
            print(f"[OK] 安装包已复制: {dest}")
        else:
            print(f"[SKIP] 安装包已存在: {dest}")


if __name__ == "__main__":
    main()
