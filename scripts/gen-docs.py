#!/usr/bin/env python3
# ──────────────────────────────────────────────────────────────────────────────
# 生成两份 Word 文档：
#   1. 发布说明（产品介绍 + 下载 + 更新内容）—— 用于在线文档发布页
#   2. 使用说明（完整操作手册）—— 从 说明文档.md 转换
#
# 用法：python scripts/gen-docs.py
# ──────────────────────────────────────────────────────────────────────────────

import json
import re
import sys
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("[FAIL] pip install python-docx")
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
PKG = json.loads((ROOT / "package.json").read_text(encoding="utf-8"))
VERSION = PKG["version"]
AUTHOR = PKG.get("author", "")
MANUAL_MD = ROOT / "说明文档.md"
OUTPUT_DIR = ROOT / "publish" / f"v{VERSION}"

BRAND_DARK = RGBColor(0x1A, 0x1A, 0x2E)
BRAND_BLUE = RGBColor(0x0D, 0x6E, 0xFD)
BRAND_ACCENT = RGBColor(0xFA, 0xB2, 0x83)
TEXT_GRAY = RGBColor(0x44, 0x44, 0x44)
TEXT_LIGHT = RGBColor(0x88, 0x88, 0x88)


def set_cell_bg(cell, color_hex):
    from lxml import etree
    tcPr = cell._element.get_or_add_tcPr()
    elm = tcPr.find(qn("w:shd"))
    if elm is None:
        elm = etree.SubElement(tcPr, qn("w:shd"))
    elm.set(qn("w:fill"), color_hex)


def setup_doc(doc):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_heading(doc, text, level=1, color=BRAND_DARK):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return h


def add_para(doc, text, size=11, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc, headers, rows, header_bg="0D6EFD"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = val
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)


# ══════════════════════════════════════════════════════════════════════════════
# 文档 1：发布说明（产品介绍页）
# ══════════════════════════════════════════════════════════════════════════════

def generate_release_page(output_path):
    doc = Document()
    setup_doc(doc)
    today = date.today().strftime("%Y-%m-%d")

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("App 冷启测速")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND_DARK
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"v{VERSION}")
    run.font.size = Pt(16)
    run.font.color.rgb = BRAND_BLUE
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Android / iOS 双平台冷启动性能测试工具")
    run.font.size = Pt(12)
    run.font.color.rgb = TEXT_GRAY
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.add_paragraph()

    # 信息表
    add_table(doc, ["属性", "内容"], [
        ("版本", f"v{VERSION}"),
        ("发布日期", today),
        ("作者", AUTHOR),
        ("适用平台", "Windows 10 / 11（64 位）"),
        ("安装包", "AppColdStart-2.0.0-setup.exe（207 MB）"),
    ])
    doc.add_paragraph()

    # 核心功能
    add_heading(doc, "核心功能", level=1)
    features = [
        ("全自动测速", "卸装 → 首次冷启动 → 杀进程 → 二次冷启动，循环 N 轮，自动去极值统计"),
        ("模板比对停表", "点选启动成功元素，运行时区域匹配，精度 ±50ms，消除人工反应误差"),
        ("scrcpy 实时镜像", "独立置顶窗口 30fps 投屏，支持后台录屏（720p/30fps）"),
        ("iOS 支持", "USB 连接 iPhone，截图 + 模板比对 + 冷启动计时"),
        ("项目持久化", "启动模板 / 跳过按钮 / 包名按项目分开存储，下次直接加载"),
        ("内置工具链", "ADB + scrcpy + libimobiledevice + Python OCR，不污染系统 PATH"),
    ]
    for title, desc in features:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{title}：")
        r.bold = True
        r.font.size = Pt(11)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r = p.add_run(desc)
        r.font.size = Pt(11)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_paragraph()

    # 安装方式
    add_heading(doc, "安装方式", level=1)
    steps = [
        "双击安装包 AppColdStart-2.0.0-setup.exe，按向导完成安装",
        "桌面 / 开始菜单出现 AppColdStart 图标",
        "双击图标启动",
        "手机用 USB 连接电脑（Android 开 USB 调试，iOS 信任电脑）",
        "在应用中选中设备，开始测速",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(step)
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    add_para(doc, "无需安装 Python，无需联网配置依赖，装完直接用。", size=11, color=BRAND_BLUE)
    doc.add_paragraph()

    # 系统要求
    add_heading(doc, "系统要求", level=1)
    add_table(doc, ["项目", "要求"], [
        ("操作系统", "Windows 10 / 11（64 位）"),
        ("Python", "无需安装（已内置）"),
        ("Android 设备", "USB 调试模式，数据线连接"),
        ("iOS 设备", "iTunes 或 AMDS 驱动，iPhone 解锁信任"),
        ("磁盘空间", "约 500 MB"),
    ])
    doc.add_paragraph()

    # 计时原理
    add_heading(doc, "计时原理", level=1)
    add_para(doc, "单一 performance.now() 时钟，不做跨进程校准：")
    add_bullet(doc, "起点：电脑发出启动命令并得到回应之后")
    add_bullet(doc, "终点：屏幕上再次出现「启动成功」模板（模板比对，精度 ±50ms）")
    add_bullet(doc, "首次、二次分开统计，样本 ≥ 3 时去一个最大一个最小")
    add_para(doc, "结论：同设备同 APK 横向对比有效；绝对值仅供参考。", size=10, color=TEXT_LIGHT)
    doc.add_paragraph()

    # 更新内容
    add_heading(doc, "v2.0.0 更新内容", level=1)
    updates = [
        ("新增", "Electron 桌面客户端，从浏览器升级为原生桌面应用"),
        ("新增", "嵌入式 Python 运行时，用户无需安装 Python"),
        ("新增", "scrcpy 实时镜像 + 后台录屏"),
        ("新增", "iOS 冷启动测试支持"),
        ("新增", "OC-2 极简主题设计"),
        ("优化", "模板比对停表精度提升至 ±50ms"),
        ("优化", "源码编译保护（.pyc），安装包不含明文代码"),
        ("修复", "路径穿越、锁不全、历史混算等 6 个高危问题"),
    ]
    for tag, desc in updates:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"[{tag}] ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = BRAND_ACCENT
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        r = p.add_run(desc)
        r.font.size = Pt(11)
        r.font.name = "Microsoft YaHei"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.add_paragraph()
    doc.add_paragraph()

    # 页脚
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"— App 冷启测速 v{VERSION} · {today} · {AUTHOR} —")
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT_LIGHT

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] 发布说明: {output_path}")


# ══════════════════════════════════════════════════════════════════════════════
# 文档 2：使用说明（从 Markdown 转换）
# ══════════════════════════════════════════════════════════════════════════════

def parse_inline(doc_paragraph, text):
    """处理 **bold** 和 `code` 等内联格式"""
    parts = re.split(r'(\*\*.+?\*\*|`.+?`)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = doc_paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = doc_paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
        elif part:
            r = doc_paragraph.add_run(part)
        if part:
            for r in doc_paragraph.runs[-1:]:
                r.font.size = Pt(11)
                r.font.name = r.font.name or "Microsoft YaHei"
                r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def md_to_docx(md_text, output_path):
    doc = Document()
    setup_doc(doc)
    lines = md_text.splitlines()
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            i += 1
            continue

        # 标题
        m = re.match(r'^(#{1,4})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level=min(level, 3))
            i += 1
            continue

        # 表格
        if line.strip().startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            headers = [c.strip() for c in line.strip().strip("|").split("|")]
            rows = []
            i += 2  # skip header + separator
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells)
                i += 1
            add_table(doc, headers, rows)
            continue

        # 引用块
        if line.strip().startswith(">"):
            text = line.strip().lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            r = p.add_run(text)
            r.font.size = Pt(10)
            r.font.color.rgb = TEXT_GRAY
            r.italic = True
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            i += 1
            continue

        # 有序列表
        m = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if m:
            p = doc.add_paragraph(style="List Number")
            parse_inline(p, m.group(2))
            i += 1
            continue

        # 无序列表
        if line.strip().startswith("- ") or line.strip().startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            parse_inline(p, line.strip()[2:])
            i += 1
            continue

        # 分隔线
        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        # 空行
        if not line.strip():
            i += 1
            continue

        # 普通段落
        p = doc.add_paragraph()
        parse_inline(p, line.strip())
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"[OK] 使用说明: {output_path}")


# ══════════════════════════════════════════════════════════════════════════════

def main():
    print(f"生成 Word 文档 (v{VERSION})...\n")

    # 1. 发布说明
    generate_release_page(OUTPUT_DIR / f"发布说明-v{VERSION}.docx")

    # 2. 使用说明（从 Markdown 转换）
    if MANUAL_MD.exists():
        md_text = MANUAL_MD.read_text(encoding="utf-8")
        md_to_docx(md_text, OUTPUT_DIR / f"使用说明-v{VERSION}.docx")
    else:
        print(f"[SKIP] {MANUAL_MD} 不存在")

    print(f"\n输出目录: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
