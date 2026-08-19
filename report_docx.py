"""性能测试报告 Word 文档生成器。

按参考文档格式一比一生成 .docx 性能报告。
参考模板：20260625-GP GOGO！Blast 1期&新4期性能对比测试报告.docx

格式要点（从参考文档解析）：
  - 字体：Microsoft YaHei（全局东亚字体）
  - Heading 1：24pt 加粗，左对齐 —— 报告标题
  - Heading 2：居中 —— 章节标题（一、二、三…）
  - Heading 3：14pt 加粗 —— 小节标题（结果汇总 / 高端机-xxx…）
  - 正文：11-12pt 加粗
  - 表格：Table Grid 样式（带边框），表头加粗居中
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# ── 全局常量 ──────────────────────────────────────────────────────────────

_FONT = "Microsoft YaHei"          # 全局东亚字体（与参考文档一致）
_COLOR_BLACK = RGBColor(0x00, 0x00, 0x00)
_COLOR_GREEN = RGBColor(0x00, 0xB1, 0x4D)   # 正向指标（优化/降低）
_COLOR_RED   = RGBColor(0xE0, 0x3E, 0x3E)   # 负向指标（劣化/增长）


# ── 辅助函数 ──────────────────────────────────────────────────────────────

def _set_font(run, size_pt: float, bold=False, color=None):
    """设置 run 字体（含东亚字体 Microsoft YaHei）。"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = _FONT
    r_el = run._element
    rPr = r_el.find(qn('w:rPr'))
    if rPr is None:
        rPr = r_el.makeelement(qn('w:rPr'), {})
        r_el.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), _FONT)
    rFonts.set(qn('w:ascii'), _FONT)
    rFonts.set(qn('w:hAnsi'), _FONT)
    if color is not None:
        run.font.color.rgb = color


def _add_para(doc, text="", size=11, bold=True, color=None,
              align=None, style=None, space_after=Pt(4)):
    """添加段落（快捷方法）。"""
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(2)
    if text:
        run = p.add_run(text)
        _set_font(run, size, bold=bold, color=color)
    return p


def _add_heading1(doc, text):
    """Heading 1：24pt 加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    _set_font(run, 24, bold=True, color=_COLOR_BLACK)
    return p


def _add_heading2(doc, text):
    """Heading 2：居中章节标题。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    _set_font(run, 16, bold=True, color=_COLOR_BLACK)
    # 底部边框线（参考文档中 Heading 2 有分隔线效果）
    _set_bottom_border(p)
    return p


def _add_heading3(doc, text):
    """Heading 3：14pt 加粗。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_font(run, 14, bold=True, color=_COLOR_BLACK)
    return p


def _set_bottom_border(paragraph):
    """给段落底部加边框线（模拟 Heading 2 下划线效果）。"""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = p_pr.makeelement(qn('w:pBdr'), {})
    bottom = p_bdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '6',
        qn('w:space'): '4',
        qn('w:color'): 'CCCCCC',
    })
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_table(doc, headers, rows, col_widths_cm=None, header_rows=1,
               data_font_size=9, header_font_size=9):
    """添加带边框的表格。

    Args:
        headers: 表头行的列表（支持多行表头，每行是一个 list）。
        rows: 数据行的列表。
        col_widths_cm: 各列宽 cm 列表（可选）。
        header_rows: 表头行数（用于合并等）。
    """
    total_cols = len(headers[0]) if isinstance(headers[0], list) else len(headers)
    # 将 headers 统一为多行格式
    if not isinstance(headers[0], list):
        headers = [headers]

    n_rows = len(headers) + len(rows)
    table = doc.add_table(rows=n_rows, cols=total_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 设置列宽
    if col_widths_cm:
        for ci, w in enumerate(col_widths_cm):
            for row in table.rows:
                row.cells[ci].width = Cm(w)

    # 填表头
    for hri, hrow in enumerate(headers):
        for ci, val in enumerate(hrow):
            cell = table.rows[hri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_font(run, header_font_size, bold=True, color=_COLOR_BLACK)
            # 表头背景色（浅灰）
            _set_cell_bg(cell, "E8EDF3")

    # 填数据
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            cell = table.rows[len(headers) + ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val) if val is not None else "/")
            _set_font(run, data_font_size, bold=True)

    return table


def _set_cell_bg(cell, color_hex):
    """设置单元格背景色。"""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): color_hex,
    })
    tc_pr.append(shd)


def _add_analysis_line(doc, label, content, label_color=None):
    """添加分析段落（标签 + 内容），如「启动时长：xxx」。

    支持多 run 着色：label 用 label_color，content 默认黑色。
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(2)
    run1 = p.add_run(label)
    _set_font(run1, 11, bold=True, color=label_color or _COLOR_BLACK)
    run2 = p.add_run(content)
    _set_font(run2, 11, bold=True, color=_COLOR_BLACK)
    return p


# ── 截尾均值（与前端 trimMean 口径一致）─────────────────────────────────

def _trim_mean(samples):
    """n>=3 去除最大最小值后取平均；否则直接平均。

    samples: list[float]（秒）
    返回: {mean, max, min, n}
    """
    if not samples:
        return {"mean": None, "max": None, "min": None, "n": 0}
    vals = list(samples)
    n = len(vals)
    if n >= 3:
        mx = max(vals)
        mn = min(vals)
        vals.remove(mx)
        vals.remove(mn)
        mean = sum(vals) / len(vals)
        return {"mean": mean, "max": mx, "min": mn, "n": n}
    mean = sum(vals) / len(vals)
    return {"mean": mean, "max": None, "min": None, "n": n}


def _fmt_sec(v):
    """秒 → 'x.xx's 字符串。"""
    if v is None:
        return "/"
    return f"{v:.2f}"


# ── 主生成函数 ────────────────────────────────────────────────────────────

def generate_report(data: dict[str, Any]) -> bytes:
    """根据前端传入的数据生成 Word 文档，返回 .docx 字节流。

    data 结构:
        title: str               报告标题
        testDate: str            测试时间
        appName: str             应用名称/包名
        platform: str            平台 (gp/ios)
        platformLabel: str       平台标签
        launchMode: str          启动方式标签
        launchModeDesc: str      启动方式描述
        plannedRounds: int       计划轮数
        totalSec: float          总耗时
        success: bool            是否成功完成
        devices: [               设备列表（每台含首/二冷启样本）
            {
                serial, model, label,
                brand, osVersion, cpu, ram, resolution,
                firsts: [ { time(ms), date, abnormal, apkVersion } ],
                seconds: [ { time(ms), date, abnormal, apkVersion } ],
            }
        ]
        auditLog: [str]          审计日志
    """
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 全局默认字体 ──
    style = doc.styles['Normal']
    style.font.name = _FONT
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), _FONT)

    devices = data.get("devices", [])
    test_date = data.get("testDate", "")
    app_name = data.get("appName", "")
    platform_label = data.get("platformLabel", "GP（安卓）")
    # iOS 首次冷启动均值调整（与前端 iosFirstAdjustSec / computeStatsGrouped 对齐）：
    # platform=='ios' 时首次均值减去此秒数（剔除 TestFlight 测试弹窗时间），GP 不调整。
    # 注意：只调展示的均值，剔除标记(MAX/MIN)仍基于原始样本，与前端 final=mean+adjust 同口径。
    platform = data.get("platform", "gp")
    _ios_adj = float(data.get("iosFirstAdjustSec", 1.0)) if platform == "ios" else 0.0

    def _first_mean(fr):
        """首次均值（iOS 时应用平台调整）；二次均值不调整。"""
        if fr["mean"] is None:
            return None
        return fr["mean"] - _ios_adj

    # 收集所有 APK 版本标签（用于多版本对比）
    apk_versions = []
    seen_ver = set()
    for dev in devices:
        for sample in dev.get("firsts", []) + dev.get("seconds", []):
            ver = sample.get("apkVersion") or "未记录版本"
            if ver not in seen_ver:
                seen_ver.add(ver)
                apk_versions.append(ver)
    if not apk_versions:
        apk_versions = ["当前版本"]

    # ═══════════════════════════════════════════════════════════════════
    # 标题（Heading 1）
    # ═══════════════════════════════════════════════════════════════════
    _add_heading1(doc, data.get("title", "冷启动测速报告"))

    # ═══════════════════════════════════════════════════════════════════
    # 一、测试背景
    # ═══════════════════════════════════════════════════════════════════
    _add_heading2(doc, "一、测试背景")
    bg_text = (
        f"本次性能测试用于验证{app_name or '应用'}在{platform_label}设备上的"
        f"冷启动性能表现，"
        f"采集首次冷启动与二次冷启动时长数据，为后续优化提供数据支撑。"
    )
    _add_para(doc, bg_text, size=12, bold=True)

    # ═══════════════════════════════════════════════════════════════════
    # 二、测试设备
    # ═══════════════════════════════════════════════════════════════════
    _add_heading2(doc, "二、测试设备")

    device_headers = ["机型", "设备编号", "厂商", "设备型号", "设备别名",
                      "系统", "CPU", "内存", "分辨率"]
    device_rows_data = []
    for i, dev in enumerate(devices):
        # 机型档位：单设备默认"测试设备"，多设备按序号
        tier = dev.get("tier") or (f"设备{i+1}" if len(devices) > 1 else "测试设备")
        device_rows_data.append([
            tier,
            dev.get("serial", "/"),
            dev.get("brand", "/"),
            dev.get("model", "/"),
            dev.get("label") or dev.get("model", "/"),
            dev.get("osVersion", "/"),
            dev.get("cpu", "/"),
            dev.get("ram", "/"),
            dev.get("resolution", "/"),
        ])
    if not device_rows_data:
        device_rows_data = [["/", "/", "/", "/", "/", "/", "/", "/", "/"]]
    _add_table(doc, device_headers, device_rows_data,
               col_widths_cm=[1.2, 1.8, 1.2, 1.8, 1.8, 1.5, 1.0, 1.0, 2.0],
               data_font_size=9, header_font_size=9)

    # ═══════════════════════════════════════════════════════════════════
    # 三、测试流程
    # ═══════════════════════════════════════════════════════════════════
    _add_heading2(doc, "三、测试流程")

    launch_desc = data.get("launchModeDesc", "模拟点击图标")
    _add_analysis_line(doc, "计时方式：",
                       f"{launch_desc}计时 ---- 屏幕进入大厅结束")
    _add_analysis_line(doc, "测试方法：",
                       f"首启/冷启时长各计 {data.get('plannedRounds', 5)} 次时长，"
                       f"去除一个最长、一个最短后对剩余样本取平均"
                       f"（样本数 < 3 时不去极值，全部取平均），"
                       f"得出最后平均启动时长")
    _add_analysis_line(doc, "平台：", platform_label)
    _add_analysis_line(doc, "启动方式：", data.get("launchMode", "模拟点击图标"))
    _add_analysis_line(doc, "启动判定：", "模板比对（cv2.matchTemplate）")
    _add_analysis_line(doc, "应用包名：", app_name or "/")
    _add_analysis_line(doc, "总耗时：", f"{data.get('totalSec', 0)}s")

    status_text = "✓ 全部完成" if data.get("success", True) else f"✗ 中止：{data.get('error', '')}"
    status_color = _COLOR_GREEN if data.get("success", True) else _COLOR_RED
    _add_analysis_line(doc, "状态：", status_text, label_color=status_color)

    # ═══════════════════════════════════════════════════════════════════
    # 四、测试结果
    # ═══════════════════════════════════════════════════════════════════
    _add_heading2(doc, "四、测试结果")

    # ── 结果汇总 ──
    _add_heading3(doc, "结果汇总")
    _add_para(doc, "性能数据包含首次启动时长、二次启动时长", size=10, bold=True)
    _add_para(doc, "  性能结论：", size=12, bold=True)

    # 计算每设备首/二均值
    dev_stats = []
    for dev in devices:
        firsts_sec = [s["time"] / 1000 for s in dev.get("firsts", [])]
        seconds_sec = [s["time"] / 1000 for s in dev.get("seconds", [])]
        fr = _trim_mean(firsts_sec)
        sr = _trim_mean(seconds_sec)
        dev_stats.append({
            "label": dev.get("label") or dev.get("model", "/"),
            "first_mean": _first_mean(fr),   # iOS 已应用平台调整
            "second_mean": sr["mean"],
            "first_n": fr["n"],
            "second_n": sr["n"],
        })

    # 多设备对比结论
    if len(dev_stats) >= 2:
        # 对比首启
        fastest_first = min(dev_stats, key=lambda x: x["first_mean"] or 999)
        slowest_first = max(dev_stats, key=lambda x: x["first_mean"] or 0)
        if fastest_first["first_mean"] and slowest_first["first_mean"]:
            diff_f = slowest_first["first_mean"] - fastest_first["first_mean"]
            pct_f = (diff_f / slowest_first["first_mean"]) * 100 if slowest_first["first_mean"] else 0
            _add_analysis_line(doc, "启动时长：",
                               f"各设备首次冷启动 {fastest_first['first_mean']:.2f}s ~ "
                               f"{slowest_first['first_mean']:.2f}s，"
                               f"最快与最慢相差 {diff_f:.2f}s（{pct_f:.1f}%）")
    elif len(dev_stats) == 1:
        ds = dev_stats[0]
        if ds["first_mean"] is not None and ds["second_mean"] is not None:
            diff = ds["first_mean"] - ds["second_mean"]
            direction = "首次比二次慢" if diff > 0 else "二次比首次慢"
            _add_analysis_line(doc, "启动时长：",
                               f"{ds['label']} 首次冷启动 {ds['first_mean']:.2f}s，"
                               f"二次冷启动 {ds['second_mean']:.2f}s，"
                               f"差值 {abs(diff):.2f}s（{direction}）")

    _add_para(doc, "")

    # ── 汇总对比表（Table 2 等效）──
    # 表头两行（与参考一致）
    t2_headers = [
        ["机型", "设备别名", "包体版本", "启动时长", "启动时长"],
        ["机型", "设备别名", "包体版本", "首启", "冷启"],
    ]
    t2_rows = []
    for i, dev in enumerate(devices):
        firsts_sec = [s["time"] / 1000 for s in dev.get("firsts", [])]
        seconds_sec = [s["time"] / 1000 for s in dev.get("seconds", [])]
        fr = _trim_mean(firsts_sec)
        sr = _trim_mean(seconds_sec)
        # 每设备每个版本一行
        ver = apk_versions[0] if apk_versions else "/"
        t2_rows.append([
            f"设备{i+1}" if len(devices) > 1 else "测试设备",
            dev.get("label") or dev.get("model", "/"),
            ver,
            _fmt_sec(_first_mean(fr)),   # iOS 已应用平台调整
            _fmt_sec(sr["mean"]),
        ])
    if not t2_rows:
        t2_rows = [["/", "/", "/", "/", "/"]]
    _add_table(doc, t2_headers, t2_rows,
               col_widths_cm=[1.5, 2.0, 2.0, 1.5, 1.5],
               data_font_size=9, header_font_size=9)

    _add_para(doc, "")

    # ═══════════════════════════════════════════════════════════════════
    # 五、首次装包启动时长
    # ═══════════════════════════════════════════════════════════════════
    _add_heading2(doc, "五、首次装包启动时长")

    _add_analysis_line(doc, "计时方式：",
                       f"{launch_desc}计时 ---- 屏幕进入大厅结束")
    _add_analysis_line(doc, "统计方法：",
                       f"首启/冷启时长各计 {data.get('plannedRounds', 5)} 次时长，"
                       f"去除一个最长、一个最短后对剩余样本取平均"
                       f"（样本数 < 3 时不去极值，全部取平均），"
                       f"得出最后平均启动时长")

    # ── 启动时长汇总表（Table 3 等效）──
    _add_para(doc, "")
    t3_headers = ["设备名称", "机型", "版本", "首次启动时长(s)", "二次冷启时长(s)"]
    t3_rows = []
    for i, dev in enumerate(devices):
        firsts_sec = [s["time"] / 1000 for s in dev.get("firsts", [])]
        seconds_sec = [s["time"] / 1000 for s in dev.get("seconds", [])]
        fr = _trim_mean(firsts_sec)
        sr = _trim_mean(seconds_sec)
        tier = f"设备{i+1}" if len(devices) > 1 else "测试设备"
        ver = apk_versions[0] if apk_versions else "/"
        t3_rows.append([
            dev.get("label") or dev.get("model", "/"),
            tier,
            ver,
            _fmt_sec(_first_mean(fr)),   # iOS 已应用平台调整
            _fmt_sec(sr["mean"]),
        ])
    if not t3_rows:
        t3_rows = [["/", "/", "/", "/", "/"]]
    _add_table(doc, t3_headers, t3_rows,
               col_widths_cm=[3.0, 1.8, 2.0, 2.5, 2.5],
               data_font_size=10, header_font_size=10)

    # ── 每设备明细 ──
    for i, dev in enumerate(devices):
        doc.add_page_break()
        label = dev.get("label") or dev.get("model", f"设备{i+1}")
        # Heading 3：设备名测试结果
        heading_text = f"{label} 测试结果"
        # 如果有 OS 信息则附上
        os_info = dev.get("osVersion")
        hardware = dev.get("hardware")
        if hardware:
            heading_text += f"（{hardware}）"
        elif os_info:
            heading_text += f"（{os_info}）"
        _add_heading3(doc, heading_text)

        # 概览对比
        _add_para(doc, "概览对比", size=12, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, color=_COLOR_BLACK)

        firsts_sec = [s["time"] / 1000 for s in dev.get("firsts", [])]
        seconds_sec = [s["time"] / 1000 for s in dev.get("seconds", [])]
        fr = _trim_mean(firsts_sec)
        sr = _trim_mean(seconds_sec)

        # 分析文本（首次均值已应用 iOS 平台调整，与前端 final 口径一致）
        fm_adj = _first_mean(fr)
        _add_analysis_line(doc, f"{label}数据：",
                           f"首次冷启动均值 {_fmt_sec(fm_adj)}s"
                           f"（{fr['n']} 样本，去极值）；"
                           f"二次冷启动均值 {_fmt_sec(sr['mean'])}s"
                           f"（{sr['n']} 样本，去极值）")
        if fm_adj is not None and sr["mean"] is not None:
            diff = fm_adj - sr["mean"]
            direction = "首次比二次慢" if diff > 0 else "首次比二次快"
            _add_analysis_line(doc, "       差值：",
                               f"{abs(diff):.2f}s（{direction}）")

        # ── 设备明细表（Table 4 等效）──
        _add_para(doc, "")
        # 多行表头（与参考文档一致）
        detail_headers = [
            ["测试版本", "启动时长", "启动时长"],
            ["测试版本", "首启", "冷启"],
        ]
        detail_rows = [[
            apk_versions[0] if apk_versions else "/",
            _fmt_sec(fm_adj),
            _fmt_sec(sr["mean"]),
        ]]
        _add_table(doc, detail_headers, detail_rows,
                   col_widths_cm=[3.0, 2.5, 2.5],
                   data_font_size=10, header_font_size=10)

        # ── 样本明细表（每条原始记录）──
        _add_para(doc, "")
        _add_para(doc, "样本明细", size=11, bold=True)

        sample_headers = ["类型", "序号", "耗时(s)", "剔除标记", "异常", "采集时间"]
        sample_rows = []

        def _fill_samples(samples, sample_type, trim_result):
            for si, s in enumerate(samples):
                tag = ""
                if trim_result["max"] is not None and s["time"] / 1000 == trim_result["max"]:
                    tag = "MAX（剔除）"
                elif trim_result["min"] is not None and s["time"] / 1000 == trim_result["min"]:
                    tag = "MIN（剔除）"
                abn = "异常" if s.get("abnormal") else ""
                sample_rows.append([
                    sample_type,
                    si + 1,
                    f"{s['time'] / 1000:.3f}",
                    tag,
                    abn,
                    s.get("date", ""),
                ])

        _fill_samples(dev.get("firsts", []), "首次冷启动", fr)
        _fill_samples(dev.get("seconds", []), "二次冷启动", sr)

        if sample_rows:
            _add_table(doc, sample_headers, sample_rows,
                       col_widths_cm=[2.0, 1.0, 1.5, 2.0, 1.0, 3.0],
                       data_font_size=9, header_font_size=9)

        # 设备名
        _add_para(doc, "")
        _add_para(doc, f"设备：{label}", size=12, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # ═══════════════════════════════════════════════════════════════════
    # 审计日志
    # ═══════════════════════════════════════════════════════════════════
    audit_log = data.get("auditLog", [])
    if audit_log:
        doc.add_page_break()
        _add_heading2(doc, "附：运行审计日志")
        for line in audit_log:
            _add_para(doc, line, size=9, bold=False,
                      align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(1))

    # ═══════════════════════════════════════════════════════════════════
    # 页脚（页码）
    # ═══════════════════════════════════════════════════════════════════
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run()
    _set_font(run, 9, bold=False)
    # 页码字段
    fld_char1 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    run._element.append(fld_char1)
    instr = run._element.makeelement(qn('w:instrText'), {})
    instr.text = 'PAGE'
    run._element.append(instr)
    fld_char2 = run._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    run._element.append(fld_char2)

    # ── 序列化 ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
